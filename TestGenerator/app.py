import random
import os
import re
from docx import Document

samples_dir = ".venv/Samples"
sample_files = [os.path.join(samples_dir, file) for file in os.listdir(samples_dir) if file.endswith(".txt")]

questions = []
for file in sample_files:
    with open(file, 'r') as f:
        content = f.read()
        questions.extend(re.findall(r'\d+\) (.+?)\n', content))

weights = []
for i in range(len(questions)):
    if 1 <= i <= 4:
        weights.append(0.3)
    elif 5 <= i <= 9:
        weights.append(0.6)
    else:
        weights.append(0.8)

def weighted_random_choice(choices, weights):
    total_weight = sum(weights)
    pick = random.uniform(0, total_weight)
    current = 0
    for choice, weight in zip(choices, weights):
        current += weight
        if pick <= current:
            return choice

selected_questions = []
selected_weights = 0
while len(selected_questions) < 14 and selected_weights < 10:
    question = weighted_random_choice(questions, weights)
    weight = weights[questions.index(question)]
    selected_questions.append(question)
    selected_weights += weight
    index = questions.index(question)
    questions.pop(index)
    weights.pop(index)

doc = Document()
doc.add_heading('Thanos Exam - ' + str (selected_weights), 0)
for i, question in enumerate(selected_questions, 1):
    doc.add_paragraph(f"{i}) {question}")

doc.save('document.docx')